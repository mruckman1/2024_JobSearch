import os
import json
import logging
import re
from datetime import datetime
from pathlib import Path
import requests
from bs4 import BeautifulSoup
import pypdf
from typing import Dict, List, Tuple
import time
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

class ResumeCustomizer:
    def __init__(self, config_path: str = "openai_config.json"):
        self.config = self._load_config(config_path)
        self.setup_folders()
        self.setup_logging()
        
        # Initialize OpenAI client
        self.client = OpenAI(api_key=self.config.get("api_key"))
        
    def _load_config(self, config_path: str) -> Dict:
        """Load configuration from JSON file."""
        default_config = {
            "model": "gpt-4o-mini-2024-07-18",
            "api_key": os.getenv("OPENAI_API_KEY", ""),
            "max_retries": 3,
            "retry_delay": 1,
            "temperature": 0.7
        }
        
        try:
            with open(config_path, 'r') as f:
                return {**default_config, **json.load(f)}
        except FileNotFoundError:
            logging.warning(f"Config file not found at {config_path}. Using defaults.")
            return default_config

    def setup_folders(self):
        """Create necessary folder structure."""
        folders = ['Base_Resume', 'Context_Files', 'Outputs']
        for folder in folders:
            Path(folder).mkdir(exist_ok=True)

    def setup_logging(self):
        """Configure logging."""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('resume_customizer.log'),
                logging.StreamHandler()
            ]
        )

    def parse_docx(self, docx_path: str) -> str:
        """Extract text content from Word document while preserving structure."""
        try:
            doc = Document(docx_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            return "\n\n".join(text_content)
        except Exception as e:
            logging.error(f"Error parsing Word document: {str(e)}")
            raise

    def scrape_job_description(self, url: str) -> str:
        """Scrape job description from URL with improved error handling."""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove unwanted elements
            for tag in ['script', 'style', 'nav', 'header', 'footer']:
                for element in soup.find_all(tag):
                    element.decompose()
            
            possible_containers = [
                soup.find(class_=class_name) for class_name in [
                    'job-description', 'description', 'posting-description',
                    'content', 'main-content', 'job-details'
                ]
            ]
            
            job_container = next((c for c in possible_containers if c), soup.find('body'))
            
            if not job_container:
                raise ValueError("Could not find job description content")
                
            text = job_container.get_text(separator='\n', strip=True)
            
            if not text:
                raise ValueError("Extracted text is empty")
                
            return text
            
        except requests.exceptions.RequestException as e:
            logging.error(f"Error fetching URL {url}: {str(e)}")
            manual_input = input("Failed to scrape job description. Would you like to paste it manually? (y/n): ")
            if manual_input.lower() == 'y':
                print("Please paste the job description (press Enter twice when finished):")
                return "\n".join(iter(input, ""))
            raise
            
        except Exception as e:
            logging.error(f"Error parsing job description: {str(e)}")
            raise

    def _query_openai(self, system_prompt: str, user_content: str = "") -> str:
        """Query OpenAI model with improved error handling."""
        retries = 0
        max_retries = self.config.get("max_retries", 3)
        
        while retries < max_retries:
            try:
                messages = [
                    {"role": "system", "content": system_prompt}
                ]
                if user_content:
                    messages.append({"role": "user", "content": user_content})
                
                response = self.client.chat.completions.create(
                    model=self.config["model"],
                    messages=messages,
                    temperature=self.config.get("temperature", 0.7)
                )
                
                return response.choices[0].message.content
                
            except Exception as e:
                retries += 1
                if retries < max_retries:
                    logging.warning(f"Attempt {retries}/{max_retries} failed: {str(e)}")
                    time.sleep(self.config.get("retry_delay", 1) * retries)
                else:
                    logging.error(f"Failed to query OpenAI after {max_retries} attempts")
                    raise

    def extract_job_requirements(self, job_description: str) -> Dict:
        """Extract key requirements from job description with enhanced keyword analysis."""
        try:
            prompt = """You are an expert ATS (Applicant Tracking System) and resume optimization specialist. 
Analyze the job description and extract key information in the following JSON format:
{
    "required_skills": [{"skill": "", "frequency": 0}],
    "required_qualifications": [{"qualification": "", "frequency": 0}],
    "key_responsibilities": [{"responsibility": "", "frequency": 0}],
    "key_terms": [{"term": "", "frequency": 0}],
    "industry_specific_keywords": [{"keyword": "", "frequency": 0}],
    "soft_skills": [{"skill": "", "frequency": 0}],
    "technical_skills": [{"skill": "", "frequency": 0}],
    "priorities": [{"priority": "", "importance_level": 0}]
}

Guidelines for extraction:
1. For each category, list items along with their frequency in the job description.
2. The "frequency" is the number of times the item appears in the job description.
3. For "priorities", assign an importance level from 1 (lowest) to 5 (highest), based on the emphasis in the job description.
4. Be thorough and detailed in your extraction. Include both explicit requirements and implicit preferences.
5. Only include items that are mentioned in the job description.

Your response should only contain the JSON data in the specified format."""
            
            response = self._query_openai(prompt, job_description)
            
            json_match = re.search(r'\{[\s\S]*\}', response)
            if not json_match:
                raise ValueError("No JSON found in response")
                
            requirements = json.loads(json_match.group())
            
            # Update expected keys to match new format
            expected_keys = [
                "required_skills", "required_qualifications", "key_responsibilities",
                "key_terms", "industry_specific_keywords", "soft_skills",
                "technical_skills", "priorities"
            ]
            if not all(key in requirements for key in expected_keys):
                raise ValueError("Missing required keys in response")
                
            return requirements
            
        except Exception as e:
            logging.error(f"Error extracting requirements: {str(e)}")
            return {
                "required_skills": [],
                "required_qualifications": [],
                "key_responsibilities": [],
                "key_terms": [],
                "industry_specific_keywords": [],
                "soft_skills": [],
                "technical_skills": [],
                "priorities": []
            }

    def customize_resume(self, resume_text: str, job_requirements: Dict) -> Tuple[str, List[Dict]]:
        """Customize resume with enhanced keyword optimization."""
        prompt = """You are an expert resume optimization specialist with deep knowledge of ATS systems. 
Your task is to enhance the resume to align with the job requirements while maintaining complete truthfulness.

First, analyze the differences between the resume and the job requirements, focusing on missing or underemphasized keywords and phrases.

Then, modify the resume to integrate the missing keywords and phrases in a natural and truthful manner.

Follow these guidelines:
1. Keyword Integration:
- Naturally incorporate missing or underemphasized key terms from the job description.
- Use industry-standard terminology where applicable.
- Ensure that the keywords are integrated in a way that accurately reflects the candidate's true experience and skills.

2. Content Optimization:
- Reorder or rephrase existing content to highlight relevant experiences and skills.
- Quantify achievements where possible.
- Use action verbs aligned with the job's terminology.

3. Authenticity Requirements:
- Never fabricate or exaggerate experiences.
- Only reword existing experience to match job terminology.
- If a required skill/experience isn't present, do not invent it.

4. Format Requirements:
- Maintain the original resume structure.
- Preserve all verifiable information (dates, titles, companies).
- Keep all modifications factual and verifiable.

Provide your response in the following format:

UPDATED_RESUME:
[Enhanced resume text here]

CHANGES:
[A JSON array of changes with:
    {
        "original": "original text",
        "updated": "updated text",
        "reason": "explanation of why this change helps match the job requirements",
        "keywords_added": ["list", "of", "relevant", "keywords", "integrated"]
    }
]
"""

        content = f"""Original Resume:
{resume_text}

Job Requirements and Keywords:
{json.dumps(job_requirements, indent=2)}

Optimization Goals:
1. Match terminology with job description where truthful
2. Emphasize relevant experiences and skills
3. Maintain professional tone and authenticity
4. Include all important job relevant keywords from the job description in the updated resume"""

        response = self._query_openai(prompt, content)
        
        try:
            if "UPDATED_RESUME:" not in response or "CHANGES:" not in response:
                raise ValueError("Response format incorrect")
                
            resume_parts = response.split("UPDATED_RESUME:")
            changes_parts = resume_parts[1].split("CHANGES:")
            
            updated_resume = changes_parts[0].strip()
            changes_text = changes_parts[1].strip()
            
            json_match = re.search(r'\[[\s\S]*\]', changes_text)
            if not json_match:
                raise ValueError("No changes JSON found in response")
                
            changes = json.loads(json_match.group())
            
            return updated_resume, changes
        except Exception as e:
            logging.error(f"Error parsing model response: {str(e)}")
            return resume_text, []

    def generate_filename(self, job_description: str, job_title: str) -> str:
        """Use LLM to generate an appropriate filename based on job details."""
        prompt = """Based on the job description and title, generate a suitable filename for a resume.
        The filename should include:
        1. The company name (extracted from the job description)
        2. The exact position title
        3. Today's date
        
        Respond with only the filename (without file extension) in the format:
        CompanyName_Position_YYYYMMDD
        
        Example: Databricks_StaffProductManager_20240316"""
        
        content = f"""Job Title: {job_title}
        
        Job Description:
        {job_description}"""
        
        try:
            filename = self._query_openai(prompt, content).strip()
            # Clean the filename of any invalid characters
            filename = "".join(c for c in filename if c.isalnum() or c in ['_', '-'])
            return filename
        except Exception as e:
            logging.error(f"Error generating filename: {str(e)}")
            # Fallback to basic filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            return f"Resume_{timestamp}"

    def save_as_docx(self, content: str, filename: str) -> str:
        """Save content as a Word document."""
        doc = Document()
        
        # Split content into paragraphs
        paragraphs = content.split('\n')
        
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                # Preserve basic formatting
                if para.isupper():  # Likely a header
                    p.runs[0].bold = True
                if len(para) < 50:  # Likely a section header
                    p.runs[0].bold = True
        
        docx_path = f"Outputs/{filename}.docx"
        doc.save(docx_path)
        return docx_path

    def convert_to_pdf(self, docx_path: str) -> str:
        """Convert Word document to PDF."""
        try:
            pdf_path = docx_path.replace('.docx', '.pdf')
            
            # Create a PDF from the Word document content
            doc = Document(docx_path)
            c = canvas.Canvas(pdf_path, pagesize=letter)
            
            # Set initial y position from top of page
            y = 750
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Handle different styles
                    if any(run.bold for run in paragraph.runs):
                        c.setFont("Helvetica-Bold", 12)
                    else:
                        c.setFont("Helvetica", 12)
                    
                    # Write text and move down
                    c.drawString(72, y, paragraph.text)
                    y -= 15
                    
                    # Check if we need a new page
                    if y < 50:
                        c.showPage()
                        y = 750
            
            c.save()
            return pdf_path
        except Exception as e:
            logging.error(f"Error converting to PDF: {str(e)}")
            raise

    def generate_output_files(self, 
                            updated_resume: str, 
                            changes: List[Dict], 
                            job_title: str,
                            job_description: str) -> Tuple[str, str, str]:
        """Generate output files with updated resume and change log."""
        try:
            # Generate filename using LLM
            base_filename = self.generate_filename(job_description, job_title)
            
            # Save as Word first
            docx_path = self.save_as_docx(updated_resume, base_filename)
            
            # Convert to PDF
            pdf_path = self.convert_to_pdf(docx_path)
            
            # Save change log
            log_filename = f"Outputs/Change_Log_{base_filename}.json"
            with open(log_filename, 'w') as f:
                json.dump(changes, f, indent=2)
            
            return docx_path, pdf_path, log_filename
        except Exception as e:
            logging.error(f"Error generating output files: {str(e)}")
            raise

def main():
    customizer = ResumeCustomizer()
    
    try:
        # Parse Word document instead of PDF
        resume_text = customizer.parse_docx("Base_Resume/resume.docx")
        logging.info("Resume parsed successfully")
        
        job_url = input("Enter job posting URL: ")
        job_description = customizer.scrape_job_description(job_url)
        logging.info("Job description scraped successfully")
        
        # Extract company and position from job description automatically
        position_info = customizer._query_openai(
            """Extract the company name and exact position title from this job description. 
            Respond in this exact format: CompanyName|PositionTitle""",
            job_description
        )
        company_name, position_title = position_info.strip().split('|')
        
        requirements = customizer.extract_job_requirements(job_description)
        logging.info("Job requirements extracted successfully")
        
        updated_resume, changes = customizer.customize_resume(resume_text, requirements)
        logging.info("Resume customization completed")
        
        # Generate all output files
        docx_path, pdf_path, log_file = customizer.generate_output_files(
            updated_resume, changes, position_title, job_description)
        
        logging.info(f"Updated resume saved to Word: {docx_path}")
        logging.info(f"Updated resume saved to PDF: {pdf_path}")
        logging.info(f"Change log saved to: {log_file}")
        
    except Exception as e:
        logging.error(f"Error in main process: {str(e)}")
        raise

if __name__ == "__main__":
    main()
