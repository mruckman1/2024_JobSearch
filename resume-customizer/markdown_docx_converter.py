# resume-customizer/markdown_docx_converter.py
from docx import Document
from docx.shared import Pt, Inches
import markdown
import html
import re

def convert_to_docx(markdown_text):
    # Create a new Document
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Clean up ampersands and other special characters before conversion
    cleaned_text = html.unescape(markdown_text)
    
    # Split the content into lines for better control
    lines = cleaned_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Handle headers
        if line.startswith('###'):
            # Remove markdown header symbols and clean the text
            header_text = line.replace('###', '').strip()
            # Unescape any remaining HTML entities
            header_text = html.unescape(header_text)
            heading = doc.add_heading(header_text, level=3)
            
        elif line.startswith('##'):
            # Remove markdown header symbols and clean the text
            header_text = line.replace('##', '').strip()
            heading = doc.add_heading(header_text, level=2)
            
        elif line.startswith('#'):
            # Remove markdown header symbols and clean the text
            header_text = line.replace('#', '').strip()
            heading = doc.add_heading(header_text, level=1)
            
        else:
            # Handle bullet points and regular text
            if line.lstrip().startswith('- ') or line.lstrip().startswith('* '):
                # This is a bullet point
                para = doc.add_paragraph(style='List Bullet')
                # Remove the bullet point character and clean the text
                text = line.lstrip('- ').lstrip('* ').strip()
                para.text = html.unescape(text)
            else:
                # Regular paragraph
                para = doc.add_paragraph()
                para.text = html.unescape(line)
            
            # Set font for the paragraph
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
    
    return doc