# modules/resume_customizer.py

import openai
import os
import logging
from docx import Document
from typing import List, Optional

def customize_resume(resume_path: str, keywords: List[str], context_docs_path: str) -> Optional[Document]:
    """Use LLM to adapt the resume, incorporating the keywords, while remaining truthful."""
    # Read the resume
    resume_text = parse_docx(resume_path)
    # Read context documents
    context_text = read_context_documents(context_docs_path)

    # Prepare the prompt
    prompt = f"""You are an expert in resume writing and optimization for Applicant Tracking Systems (ATS). Given the candidate's resume and additional context documents, update the resume to include the following keywords, ensuring that the resume remains truthful and accurately reflects the candidate's experience.

Keywords:
{', '.join(keywords)}

Resume:
{resume_text}

Context Documents:
{context_text}

Updated Resume:"""

    # Use OpenAI API to generate the updated resume
    openai.api_key = os.getenv("OPENAI_API_KEY")
    try:
        response = openai.Completion.create(
            engine="text-davinci-003",
            prompt=prompt,
            max_tokens=2048,
            temperature=0.5,
        )
        updated_resume_text = response.choices[0].text.strip()
        # Convert updated_resume_text to a Document object
        updated_doc = Document()
        for paragraph in updated_resume_text.split('\n'):
            if paragraph.strip() != '':
                updated_doc.add_paragraph(paragraph)
        return updated_doc
    except Exception as e:
        logging.error(f"Error customizing resume: {e}")
        return None

def parse_docx(docx_path: str) -> str:
    """Extract text from a .docx file."""
    doc = Document(docx_path)
    text = '\n'.join([para.text for para in doc.paragraphs])
    return text

def read_context_documents(context_docs_path: str) -> str:
    """Read and concatenate all context documents."""
    context_text = ""
    for filename in os.listdir(context_docs_path):
        file_path = os.path.join(context_docs_path, filename)
        if filename.endswith('.txt'):
            with open(file_path, 'r') as f:
                context_text += f.read() + '\n'
        elif filename.endswith('.docx'):
            context_text += parse_docx(file_path) + '\n'
    return context_text
